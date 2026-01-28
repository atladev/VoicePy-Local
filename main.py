import os
import time
from pathlib import Path
from docx import Document
import torch
from shutil import copy2, rmtree
from TTS.api import TTS
from TTS.utils.synthesizer import Synthesizer
import io
from contextlib import redirect_stdout
import streamlit as st
import psutil
import re
from datetime import datetime

# =============================
# CONSOLE STATUS / LOG HELPERS
# =============================
_last_standby_print = 0.0

class TerminalDisplay:
    """Classe para exibir informações formatadas no terminal"""
    
    @staticmethod
    def clear_screen():
        """Limpa a tela do terminal"""
        os.system('cls' if os.name == 'nt' else 'clear')
    
    @staticmethod
    def print_header():
        """Imprime o cabeçalho do painel"""
        print("\n" + "="*80)
        print("█" * 80)
        print(f"{'DEDÉ LABS® - SISTEMA DE LOCUÇÃO':^80}")
        print(f"{'Painel de Controle v1.04':^80}")
        print("█" * 80)
        print("="*80 + "\n")
    
    @staticmethod
    def print_status_box(status: str, color_code: str = "37"):
        """Imprime uma caixa de status colorida"""
        status_map = {
            "LIVRE": ("✅ SISTEMA LIVRE - Pode usar!", "42"),  # Verde
            "EM_USO": ("🚫 SISTEMA EM USO - Alguém está gerando áudio!", "41"),  # Vermelho
            "STANDBY": ("⏳ SISTEMA EM STANDBY - Aguardando ação...", "43"),  # Amarelo
            "PROCESSANDO": ("⚙️  PROCESSANDO - Gerando áudios...", "44"),  # Azul
        }
        
        msg, bg = status_map.get(status, (status, color_code))
        
        print("\n┌" + "─"*78 + "┐")
        print(f"│\033[{bg};97m{msg:^78}\033[0m│")
        print("└" + "─"*78 + "┘\n")
    
    @staticmethod
    def print_info_table(info_dict: dict):
        """Imprime tabela de informações"""
        print("┌" + "─"*78 + "┐")
        print(f"│ {'INFORMAÇÕES DO SISTEMA':^76} │")
        print("├" + "─"*78 + "┤")
        
        for key, value in info_dict.items():
            key_str = f"{key}:"
            print(f"│ {key_str:<30} {str(value):<45} │")
        
        print("└" + "─"*78 + "┘\n")
    
    @staticmethod
    def print_progress_bar(current: int, total: int, prefix: str = "Progresso"):
        """Imprime barra de progresso"""
        percent = (current / total) * 100 if total > 0 else 0
        filled = int(50 * current / total) if total > 0 else 0
        bar = "█" * filled + "░" * (50 - filled)
        
        print(f"\r{prefix}: |{bar}| {percent:.1f}% ({current}/{total})", end="", flush=True)
        
        if current == total:
            print()  # Nova linha ao completar

terminal = TerminalDisplay()

def _console(status: str, msg: str = ""):
    """Log melhorado com timestamp e formatação"""
    ts = time.strftime('%Y-%m-%d %H:%M:%S')
    
    # Mapeia cores para diferentes tipos de status
    color_map = {
        "INFO": "\033[36m",      # Ciano
        "EXECUTANDO": "\033[35m", # Magenta
        "SUCESSO": "\033[32m",    # Verde
        "ERRO": "\033[31m",       # Vermelho
        "AVISO": "\033[33m",      # Amarelo
        "STANDBY": "\033[37m",    # Branco
        "FLUSH": "\033[34m",      # Azul
    }
    
    color = color_map.get(status.upper(), "\033[37m")
    reset = "\033[0m"
    
    # Símbolos para cada tipo
    symbol_map = {
        "INFO": "ℹ️ ",
        "EXECUTANDO": "▶️ ",
        "SUCESSO": "✅",
        "ERRO": "❌",
        "AVISO": "⚠️ ",
        "STANDBY": "💤",
        "FLUSH": "🧹",
    }
    
    symbol = symbol_map.get(status.upper(), "•")
    
    try:
        line = f"{color}[{ts}] {symbol} [{status.upper():^12}]{reset} {msg}"
        print(line, flush=True)
    except Exception:
        pass

def _console_standby_throttled(msg: str, interval: float = 8.0):
    global _last_standby_print
    now = time.time()
    if now - _last_standby_print >= interval:
        _console("STANDBY", msg)
        _last_standby_print = now

def update_terminal_display(status: str = "STANDBY", extra_info: dict = None):
    """Atualiza o display completo do terminal"""
    terminal.clear_screen()
    terminal.print_header()
    terminal.print_status_box(status)
    
    # Informações básicas do sistema
    info = {
        "Data/Hora": datetime.now().strftime("%d/%m/%Y %H:%M:%S"),
        "Modelo TTS": params.get("model_name", "N/A"),
        "Dispositivo": params.get("device", "N/A"),
        "Idioma": params.get("language", "N/A"),
        "Status do Lock": "🔒 Bloqueado" if is_app_in_use() else "🔓 Livre",
    }
    
    if extra_info:
        info.update(extra_info)
    
    terminal.print_info_table(info)
    
    # Log de atividades recentes
    print("📋 ÚLTIMAS ATIVIDADES:")
    print("─" * 80)

# =============================
# CONFIGURAÇÕES GERAIS
# =============================
p = psutil.Process(os.getpid())
try:
    p.nice(psutil.HIGH_PRIORITY_CLASS)
except Exception:
    pass

os.environ["COQUI_TOS_AGREED"] = "1"

DEFAULT_MODEL = "tts_models/multilingual/multi-dataset/xtts_v2"
DEFAULT_DEVICE = "cuda" if torch.cuda.is_available() else "cpu"

DOWNLOAD_PATH = r"C:\\Users\\dud\\Downloads\\youtube\\VIAJENS\\AUDIOS_FEITOS"
textos_com_erro = []

# =============================
# SISTEMA DE BLOQUEIO GLOBAL
# =============================
LOCK_FILE = Path("app_in_use.lock")

def set_app_status(in_use: bool):
    """Cria ou remove o arquivo de bloqueio para sinalizar uso."""
    if in_use:
        LOCK_FILE.write_text("IN_USE")
        _console("INFO", "Sistema bloqueado para uso exclusivo")
        update_terminal_display("EM_USO")
    else:
        if LOCK_FILE.exists():
            LOCK_FILE.unlink(missing_ok=True)
        _console("INFO", "Sistema liberado para uso")
        update_terminal_display("LIVRE")

def is_app_in_use() -> bool:
    """Verifica se o sistema está em uso por outro usuário."""
    return LOCK_FILE.exists()

def cleanup_stale_lock():
    """Remove lock file antigo na inicialização"""
    if LOCK_FILE.exists():
        try:
            # Remove o lock na inicialização
            LOCK_FILE.unlink(missing_ok=True)
            _console("INFO", "Lock antigo removido na inicialização")
        except Exception as e:
            _console("ERRO", f"Erro ao remover lock: {e}")

# =============================
# UTILITÁRIOS
# =============================
def sanitize_name(name: str) -> str:
    base = re.sub(r"[\\/:*?\"<>|]", "_", name)
    base = re.sub(r"\s+", " ", base).strip()
    return base

def list_wav_files(folder: str):
    try:
        p = Path(folder)
        if not p.exists():
            return []
        return [str(f.name) for f in p.glob("*.wav")]
    except Exception:
        return []

params = {
    "remove_trailing_dots": True,
    "voice": "",
    "language": "pt",
    "model_name": DEFAULT_MODEL,
    "device": DEFAULT_DEVICE,
}

def new_split_into_sentences(self, text):
    sentences = self.seg.segment(text)
    if params['remove_trailing_dots']:
        sentences_without_dots = []
        for sentence in sentences:
            if sentence.endswith('.') and not sentence.endswith('...'):
                sentence = sentence[:-1]
            sentences_without_dots.append(sentence)
        return sentences_without_dots
    else:
        return sentences

Synthesizer.split_into_sentences = new_split_into_sentences

# =============================
# CARREGAMENTO DO MODELO
# =============================
@st.cache_resource(show_spinner=True)
def load_model(model_name: str, device: str):
    _console("INFO", f"Carregando modelo: {model_name}")
    return TTS(model_name).to(device)

def flush_tts_cache():
    """Força descarregar o modelo atual do cache e da GPU."""
    try:
        _console("FLUSH", "Limpando cache do modelo TTS...")
        st.cache_resource.clear()
        torch.cuda.empty_cache()
        _console("SUCESSO", "Cache limpo com sucesso!")
    except Exception as e:
        _console("ERRO", f"Falha ao limpar cache: {e}")

# =============================
# LEITURA DO DOCX
# =============================
@st.cache_data(show_spinner=False)
def load_text(file_path):
    doc = Document(file_path)
    return [paragraph.text.replace('.', ',.') for paragraph in doc.paragraphs if paragraph.text.strip()]

def generate_audio_filename(index):
    return f"audio_{index + 1}.wav"

# =============================
# GERAÇÃO COM LOG
# =============================
def tts_to_file_logged(model, text: str, out_path: str, language: str, speaker_wav: str, speed: float = 0.85):
    output_buffer = io.StringIO()
    with redirect_stdout(output_buffer):
        model.tts_to_file(
            text=text,
            file_path=out_path,
            speaker_wav=speaker_wav,
            language=language,
            speed=speed,
        )
    return output_buffer.getvalue()

# =============================
# APP STREAMLIT
# =============================
def main_app():
    # LIMPEZA DO LOCK NA INICIALIZAÇÃO
    cleanup_stale_lock()
    
    # Inicialização do display no terminal
    update_terminal_display("STANDBY", {
        "GPU Disponível": "✅ Sim" if torch.cuda.is_available() else "❌ Não",
        "Memória GPU": f"{torch.cuda.get_device_properties(0).total_memory / 1e9:.1f} GB" if torch.cuda.is_available() else "N/A"
    })
    
    _console("INFO", "Aplicação Streamlit iniciada")
    
    st.set_page_config(page_title="Dedé Labs® -------- Créditos gratis: 2000", layout="wide")

    st.title("Dedé Labs® -------- Créditos gratis: 2000")
    st.caption("versão 1.04")

    # ===================================
    # AVISO DE USO ATIVO (BANNER GLOBAL)
    # ===================================
    app_status_placeholder = st.empty()
    # sempre começa como liberado
    app_status_placeholder.markdown(
        "<div style='background-color:#4CAF50;padding:10px;border-radius:8px;text-align:center;color:white;font-weight:bold;'>✅ Pode usar! Ninguém está usando agora.</div>",
        unsafe_allow_html=True
    )

    with st.sidebar:
        st.subheader("Configurações")

        language = st.radio("Idioma da narração", options=["pt", "es"], index=0, horizontal=True)
        params["language"] = language

        default_voice_dir = r"C:\\Users\\dud\\Downloads\\youtube\\scripts\\voices"
        voice_dir = st.text_input("Pasta com vozes (.wav)", value=default_voice_dir)
        wavs = list_wav_files(voice_dir)

        if not wavs:
            st.warning("Nenhum arquivo .wav encontrado nessa pasta.")
        else:
            selected_wav_name = st.selectbox("Escolha a voz (arquivo .wav)", wavs)
            selected_wav_path = str(Path(voice_dir) / selected_wav_name)

            # Se a voz mudou, força flush do modelo
            if params.get("voice") and params["voice"] != selected_wav_path:
                flush_tts_cache()

            params["voice"] = selected_wav_path
            _console("INFO", f"Voz selecionada: {selected_wav_name}")

            st.markdown("**Pré-escuta da voz selecionada**")
            try:
                with open(selected_wav_path, "rb") as f:
                    st.audio(f.read(), format="audio/wav")
            except Exception as e:
                st.error(f"Erro ao carregar preview da voz: {e}")

        with st.expander("Opções avançadas"):
            params["model_name"] = st.text_input("Modelo Coqui", value=params["model_name"])
            params["device"] = st.selectbox("Dispositivo", ["cuda", "cpu"], index=0 if DEFAULT_DEVICE == "cuda" else 1)
            params["remove_trailing_dots"] = st.checkbox("Remover ponto final simples de frases", value=True)

    st.divider()

    st.markdown("### Teste rápido da voz selecionada (TTS)")
    sample_text = st.text_input("Texto de teste", value="Este é um teste de narração.")
    if st.button("Reproduzir amostra gerada"):
        if not params.get("voice"):
            st.error("Selecione uma voz primeiro.")
        elif is_app_in_use():
            st.error("Outro usuário está usando agora. Aguarde o aviso verde.")
        else:
            set_app_status(True)
            app_status_placeholder.markdown(
                "<div style='background-color:#ff4d4d;padding:10px;border-radius:8px;text-align:center;color:white;font-weight:bold;'>🚫 Tem gente usando! Não mexa em nada agora.</div>",
                unsafe_allow_html=True
            )
            try:
                _console("EXECUTANDO", "Gerando amostra de teste...")
                with st.spinner("Gerando amostra..."):
                    model = load_model(params["model_name"], params["device"])
                    tmp_sample = Path("./_tmp_sample.wav")
                    _ = tts_to_file_logged(
                        model,
                        text=sample_text,
                        out_path=str(tmp_sample),
                        language=params["language"],
                        speaker_wav=params["voice"],
                        speed=0.9,
                    )
                    with open(tmp_sample, "rb") as f:
                        st.audio(f.read(), format="audio/wav")
                    _console("SUCESSO", "Amostra gerada com sucesso!")
            finally:
                if tmp_sample.exists():
                    tmp_sample.unlink(missing_ok=True)
                set_app_status(False)
                app_status_placeholder.markdown(
                    "<div style='background-color:#4CAF50;padding:10px;border-radius:8px;text-align:center;color:white;font-weight:bold;'>✅ Pode usar! Ninguém está usando agora.</div>",
                    unsafe_allow_html=True
                )

    st.divider()

    uploaded_file = st.file_uploader("Selecione o arquivo .docx", type="docx")
    if uploaded_file is not None:
        if is_app_in_use():
            st.error("Outro usuário está gerando áudio agora. Aguarde o aviso verde.")
            return

        set_app_status(True)
        app_status_placeholder.markdown(
            "<div style='background-color:#ff4d4d;padding:10px;border-radius:8px;text-align:center;color:white;font-weight:bold;'>🚫 Tem gente usando! Não mexa em nada agora.</div>",
            unsafe_allow_html=True
        )

        try:
            _console("EXECUTANDO", f"Arquivo recebido: {uploaded_file.name}")
            original_name = sanitize_name(uploaded_file.name)
            base_name = os.path.splitext(original_name)[0]
            new_folder_name = f"{params['language']}_{base_name}"
            new_folder_path = os.path.join(DOWNLOAD_PATH, new_folder_name)
            os.makedirs(new_folder_path, exist_ok=True)
            
            _console("INFO", f"Pasta de destino: {new_folder_path}")

            docx_destination = os.path.join(new_folder_path, original_name)
            with open(docx_destination, "wb") as f:
                f.write(uploaded_file.read())

            paragraphs = load_text(docx_destination)
            _console("INFO", f"Total de parágrafos a processar: {len(paragraphs)}")
            
            update_terminal_display("PROCESSANDO", {
                "Arquivo": uploaded_file.name,
                "Total de parágrafos": len(paragraphs),
                "Pasta destino": new_folder_name
            })
            
            model = load_model(params["model_name"], params["device"])

            st.write("Iniciando geração dos áudios...")
            progress = st.progress(0)
            log_placeholder = st.empty()
            textos_com_erro.clear()
            error_texts = []

            for index, paragraph in enumerate(paragraphs):
                output_file_name = generate_audio_filename(index)
                output_file_path = os.path.join(new_folder_path, output_file_name)

                st.write(f"Gerando áudio para o parágrafo {index + 1}/{len(paragraphs)}...")
                _console("EXECUTANDO", f"Processando parágrafo {index + 1}/{len(paragraphs)}")
                
                # Atualiza barra de progresso no terminal
                terminal.print_progress_bar(index + 1, len(paragraphs), "Geração de áudios")
                
                try:
                    log = tts_to_file_logged(
                        model,
                        text=paragraph,
                        out_path=output_file_path,
                        language=params["language"],
                        speaker_wav=params["voice"],
                        speed=0.85,
                    )
                    log_placeholder.code(log or "(sem logs)")
                    _console("SUCESSO", f"Áudio {index + 1} gerado: {output_file_name}")
                except Exception as e:
                    log_placeholder.error(f"Erro ao gerar áudio: {e}")
                    _console("ERRO", f"Falha no parágrafo {index + 1}: {str(e)[:100]}")
                    if "exceeds the character limit" in str(e).lower():
                        error_texts.append(paragraph)
                    continue

                if "exceeds the character limit" in (log or "").lower():
                    novo_nome = f"audio_{index + 1}__pode ter erro.wav"
                    os.rename(output_file_path, os.path.join(new_folder_path, novo_nome))
                    textos_com_erro.append(paragraph)
                    _console("AVISO", f"Parágrafo {index + 1} pode conter erros")

                progress.progress((index + 1) / len(paragraphs))

            if error_texts or textos_com_erro:
                error_docx_path = os.path.join(new_folder_path, "paragrafos_com_erro.docx")
                error_doc = Document()
                for t in error_texts + textos_com_erro:
                    error_doc.add_paragraph(t)
                error_doc.save(error_docx_path)
                st.warning(f"Parágrafos com possível erro salvos em: {error_docx_path}")
                _console("AVISO", f"{len(error_texts) + len(textos_com_erro)} parágrafos com erro salvos")

            st.success("Processo de geração concluído!")
            _console("SUCESSO", "🎉 Processo completo! Todos os áudios foram gerados.")
            
        finally:
            set_app_status(False)
            app_status_placeholder.markdown(
                "<div style='background-color:#4CAF50;padding:10px;border-radius:8px;text-align:center;color:white;font-weight:bold;'>✅ Pode usar! Ninguém está usando agora.</div>",
                unsafe_allow_html=True
            )

    if uploaded_file is None and not is_app_in_use():
        _console_standby_throttled("💤 Sistema em standby - Aguardando ação do usuário.")

if __name__ == "__main__":
    main_app()
