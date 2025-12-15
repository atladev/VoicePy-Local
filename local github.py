"""
Modern Text-to-Speech Audio Generator
A sleek desktop application for generating audio from text documents using AI voices.
"""

import os
import re
import time
import threading
from pathlib import Path
from tkinter import *
from tkinter import ttk, filedialog, messagebox
from docx import Document
import torch
from TTS.api import TTS
from TTS.utils.synthesizer import Synthesizer
import pygame
from PIL import Image, ImageTk

# =============================
# CONFIGURATION
# =============================
os.environ["COQUI_TOS_AGREED"] = "1"

DEFAULT_MODEL = "tts_models/multilingual/multi-dataset/xtts_v2"
DEFAULT_DEVICE = "cuda" if torch.cuda.is_available() else "cpu"
DEFAULT_OUTPUT_DIR = str(Path.home() / "TTS_Output")

# =============================
# UTILITIES
# =============================
def sanitize_name(name: str) -> str:
    """Clean filename from invalid characters."""
    base = re.sub(r'[\\/:*?"<>|]', "_", name)
    return re.sub(r"\s+", " ", base).strip()

def list_wav_files(folder: str):
    """List all .wav files in a folder."""
    try:
        p = Path(folder)
        if not p.exists():
            return []
        return sorted([str(f) for f in p.glob("*.wav")])
    except Exception:
        return []

# Patch for sentence splitting
def new_split_into_sentences(self, text):
    sentences = self.seg.segment(text)
    return [s[:-1] if s.endswith('.') and not s.endswith('...') else s for s in sentences]

Synthesizer.split_into_sentences = new_split_into_sentences

# =============================
# MODERN GUI APPLICATION
# =============================
class ModernTTSApp:
    def __init__(self, root):
        self.root = root
        self.root.title("TTS Audio Generator")
        self.root.geometry("900x700")
        self.root.minsize(800, 600)
        
        # Variables
        self.model = None
        self.is_generating = False
        self.selected_voice = None
        self.selected_docx = None
        self.output_dir = DEFAULT_OUTPUT_DIR
        
        # Colors - Modern dark theme
        self.bg_dark = "#1e1e2e"
        self.bg_medium = "#2a2a3e"
        self.bg_light = "#363650"
        self.accent = "#6c63ff"
        self.accent_hover = "#5a52d5"
        self.text_color = "#e0e0e0"
        self.success = "#00d4aa"
        self.warning = "#ffd93d"
        self.error = "#ff6b6b"
        
        self.root.configure(bg=self.bg_dark)
        
        # Initialize pygame for audio playback
        pygame.mixer.init()
        
        self.setup_ui()
        self.create_output_dir()
        
    def create_output_dir(self):
        """Ensure output directory exists."""
        Path(self.output_dir).mkdir(parents=True, exist_ok=True)
        
    def setup_ui(self):
        """Create the modern UI layout."""
        # Header
        header = Frame(self.root, bg=self.bg_medium, height=80)
        header.pack(fill=X, pady=(0, 20))
        header.pack_propagate(False)
        
        title = Label(
            header, 
            text="🎙️ TTS Audio Generator",
            font=("Segoe UI", 24, "bold"),
            bg=self.bg_medium,
            fg=self.text_color
        )
        title.pack(pady=20)
        
        # Main container
        main = Frame(self.root, bg=self.bg_dark)
        main.pack(fill=BOTH, expand=True, padx=30, pady=(0, 20))
        
        # Left panel - Voice selection
        left_panel = self.create_card(main, "Voice Settings")
        left_panel.pack(side=LEFT, fill=BOTH, expand=True, padx=(0, 10))
        
        self.create_voice_section(left_panel)
        
        # Right panel - Document processing
        right_panel = self.create_card(main, "Document Processing")
        right_panel.pack(side=RIGHT, fill=BOTH, expand=True, padx=(10, 0))
        
        self.create_document_section(right_panel)
        
        # Bottom panel - Progress and status
        bottom_panel = Frame(self.root, bg=self.bg_dark)
        bottom_panel.pack(fill=X, padx=30, pady=(0, 20))
        
        self.create_progress_section(bottom_panel)
        
    def create_card(self, parent, title):
        """Create a modern card container."""
        card = Frame(parent, bg=self.bg_medium, relief=FLAT)
        card.pack(fill=BOTH, expand=True)
        
        card_title = Label(
            card,
            text=title,
            font=("Segoe UI", 14, "bold"),
            bg=self.bg_medium,
            fg=self.text_color,
            anchor=W
        )
        card_title.pack(fill=X, padx=20, pady=(15, 10))
        
        return card
        
    def create_voice_section(self, parent):
        """Create voice selection UI."""
        content = Frame(parent, bg=self.bg_medium)
        content.pack(fill=BOTH, expand=True, padx=20, pady=(0, 20))
        
        # Voice folder selection
        Label(
            content,
            text="Voice Folder:",
            font=("Segoe UI", 10),
            bg=self.bg_medium,
            fg=self.text_color,
            anchor=W
        ).pack(fill=X, pady=(10, 5))
        
        folder_frame = Frame(content, bg=self.bg_medium)
        folder_frame.pack(fill=X, pady=(0, 15))
        
        self.voice_folder_entry = Entry(
            folder_frame,
            font=("Segoe UI", 10),
            bg=self.bg_light,
            fg=self.text_color,
            insertbackground=self.text_color,
            relief=FLAT
        )
        self.voice_folder_entry.pack(side=LEFT, fill=X, expand=True, ipady=8)
        self.voice_folder_entry.insert(0, str(Path.home()))
        
        browse_btn = self.create_button(
            folder_frame,
            "Browse",
            lambda: self.browse_voice_folder(),
            width=10
        )
        browse_btn.pack(side=RIGHT, padx=(10, 0))
        
        # Voice list
        Label(
            content,
            text="Available Voices:",
            font=("Segoe UI", 10),
            bg=self.bg_medium,
            fg=self.text_color,
            anchor=W
        ).pack(fill=X, pady=(0, 5))
        
        list_frame = Frame(content, bg=self.bg_light)
        list_frame.pack(fill=BOTH, expand=True, pady=(0, 15))
        
        scrollbar = Scrollbar(list_frame)
        scrollbar.pack(side=RIGHT, fill=Y)
        
        self.voice_listbox = Listbox(
            list_frame,
            font=("Segoe UI", 9),
            bg=self.bg_light,
            fg=self.text_color,
            selectbackground=self.accent,
            selectforeground="white",
            relief=FLAT,
            yscrollcommand=scrollbar.set
        )
        self.voice_listbox.pack(fill=BOTH, expand=True, padx=5, pady=5)
        scrollbar.config(command=self.voice_listbox.yview)
        
        # Preview button
        preview_btn = self.create_button(
            content,
            "🔊 Preview Selected Voice",
            self.preview_voice,
            full_width=True
        )
        preview_btn.pack(fill=X)
        
        # Language selection
        Label(
            content,
            text="Language:",
            font=("Segoe UI", 10),
            bg=self.bg_medium,
            fg=self.text_color,
            anchor=W
        ).pack(fill=X, pady=(15, 5))
        
        self.language_var = StringVar(value="en")
        lang_frame = Frame(content, bg=self.bg_medium)
        lang_frame.pack(fill=X)
        
        for lang, label in [("en", "English"), ("pt", "Portuguese"), ("es", "Spanish")]:
            Radiobutton(
                lang_frame,
                text=label,
                variable=self.language_var,
                value=lang,
                font=("Segoe UI", 9),
                bg=self.bg_medium,
                fg=self.text_color,
                selectcolor=self.bg_light,
                activebackground=self.bg_medium,
                activeforeground=self.text_color
            ).pack(side=LEFT, padx=(0, 15))
        
    def create_document_section(self, parent):
        """Create document processing UI."""
        content = Frame(parent, bg=self.bg_medium)
        content.pack(fill=BOTH, expand=True, padx=20, pady=(0, 20))
        
        # Document selection
        Label(
            content,
            text="Document (.docx):",
            font=("Segoe UI", 10),
            bg=self.bg_medium,
            fg=self.text_color,
            anchor=W
        ).pack(fill=X, pady=(10, 5))
        
        self.doc_label = Label(
            content,
            text="No document selected",
            font=("Segoe UI", 9),
            bg=self.bg_light,
            fg="#888",
            anchor=W,
            padx=15,
            pady=15,
            relief=FLAT
        )
        self.doc_label.pack(fill=X, pady=(0, 10))
        
        select_doc_btn = self.create_button(
            content,
            "📄 Select Document",
            self.select_document,
            full_width=True
        )
        select_doc_btn.pack(fill=X, pady=(0, 15))
        
        # Output directory
        Label(
            content,
            text="Output Directory:",
            font=("Segoe UI", 10),
            bg=self.bg_medium,
            fg=self.text_color,
            anchor=W
        ).pack(fill=X, pady=(0, 5))
        
        output_frame = Frame(content, bg=self.bg_medium)
        output_frame.pack(fill=X, pady=(0, 20))
        
        self.output_entry = Entry(
            output_frame,
            font=("Segoe UI", 10),
            bg=self.bg_light,
            fg=self.text_color,
            insertbackground=self.text_color,
            relief=FLAT
        )
        self.output_entry.pack(side=LEFT, fill=X, expand=True, ipady=8)
        self.output_entry.insert(0, self.output_dir)
        
        browse_output_btn = self.create_button(
            output_frame,
            "Browse",
            self.browse_output_dir,
            width=10
        )
        browse_output_btn.pack(side=RIGHT, padx=(10, 0))
        
        # Test TTS
        Label(
            content,
            text="Test Text:",
            font=("Segoe UI", 10),
            bg=self.bg_medium,
            fg=self.text_color,
            anchor=W
        ).pack(fill=X, pady=(0, 5))
        
        self.test_text = Text(
            content,
            font=("Segoe UI", 9),
            bg=self.bg_light,
            fg=self.text_color,
            insertbackground=self.text_color,
            relief=FLAT,
            height=4,
            wrap=WORD
        )
        self.test_text.pack(fill=X, pady=(0, 10))
        self.test_text.insert("1.0", "This is a test of the text to speech system.")
        
        test_btn = self.create_button(
            content,
            "🎵 Generate Test Audio",
            self.test_tts,
            full_width=True
        )
        test_btn.pack(fill=X, pady=(0, 20))
        
        # Generate button
        self.generate_btn = self.create_button(
            content,
            "⚡ Generate All Audio Files",
            self.generate_audio,
            full_width=True,
            accent=True
        )
        self.generate_btn.pack(fill=X)
        
    def create_progress_section(self, parent):
        """Create progress and status UI."""
        # Status label
        self.status_label = Label(
            parent,
            text="Ready",
            font=("Segoe UI", 10),
            bg=self.bg_dark,
            fg=self.text_color,
            anchor=W
        )
        self.status_label.pack(fill=X, pady=(0, 5))
        
        # Progress bar
        style = ttk.Style()
        style.theme_use('clam')
        style.configure(
            "Custom.Horizontal.TProgressbar",
            background=self.accent,
            troughcolor=self.bg_medium,
            borderwidth=0,
            lightcolor=self.accent,
            darkcolor=self.accent
        )
        
        self.progress = ttk.Progressbar(
            parent,
            style="Custom.Horizontal.TProgressbar",
            mode='determinate'
        )
        self.progress.pack(fill=X)
        
    def create_button(self, parent, text, command, width=None, full_width=False, accent=False):
        """Create a modern styled button."""
        bg = self.accent if accent else self.bg_light
        hover_bg = self.accent_hover if accent else "#464660"
        
        btn = Button(
            parent,
            text=text,
            command=command,
            font=("Segoe UI", 10, "bold" if accent else "normal"),
            bg=bg,
            fg="white",
            activebackground=hover_bg,
            activeforeground="white",
            relief=FLAT,
            cursor="hand2",
            padx=20,
            pady=10
        )
        
        if width:
            btn.config(width=width)
            
        # Hover effects
        btn.bind("<Enter>", lambda e: btn.config(bg=hover_bg))
        btn.bind("<Leave>", lambda e: btn.config(bg=bg))
        
        return btn
        
    def browse_voice_folder(self):
        """Browse for voice folder."""
        folder = filedialog.askdirectory(title="Select Voice Folder")
        if folder:
            self.voice_folder_entry.delete(0, END)
            self.voice_folder_entry.insert(0, folder)
            self.load_voices()
            
    def load_voices(self):
        """Load available voice files."""
        folder = self.voice_folder_entry.get()
        voices = list_wav_files(folder)
        
        self.voice_listbox.delete(0, END)
        for voice in voices:
            self.voice_listbox.insert(END, Path(voice).name)
            
    def preview_voice(self):
        """Preview selected voice file."""
        selection = self.voice_listbox.curselection()
        if not selection:
            messagebox.showwarning("No Selection", "Please select a voice first.")
            return
            
        folder = self.voice_folder_entry.get()
        voice_name = self.voice_listbox.get(selection[0])
        voice_path = str(Path(folder) / voice_name)
        
        try:
            pygame.mixer.music.load(voice_path)
            pygame.mixer.music.play()
            self.status_label.config(text=f"Playing: {voice_name}")
        except Exception as e:
            messagebox.showerror("Error", f"Could not play voice: {e}")
            
    def select_document(self):
        """Select document file."""
        file = filedialog.askopenfilename(
            title="Select Document",
            filetypes=[("Word Documents", "*.docx")]
        )
        if file:
            self.selected_docx = file
            self.doc_label.config(
                text=Path(file).name,
                fg=self.success
            )
            
    def browse_output_dir(self):
        """Browse for output directory."""
        folder = filedialog.askdirectory(title="Select Output Directory")
        if folder:
            self.output_dir = folder
            self.output_entry.delete(0, END)
            self.output_entry.insert(0, folder)
            
    def load_tts_model(self):
        """Load TTS model (cached)."""
        if self.model is None:
            self.status_label.config(text="Loading TTS model...")
            self.root.update()
            self.model = TTS(DEFAULT_MODEL).to(DEFAULT_DEVICE)
            self.status_label.config(text="Model loaded successfully")
            
    def test_tts(self):
        """Test TTS with sample text."""
        selection = self.voice_listbox.curselection()
        if not selection:
            messagebox.showwarning("No Voice", "Please select a voice first.")
            return
            
        text = self.test_text.get("1.0", END).strip()
        if not text:
            messagebox.showwarning("No Text", "Please enter test text.")
            return
            
        def generate():
            try:
                self.load_tts_model()
                
                folder = self.voice_folder_entry.get()
                voice_name = self.voice_listbox.get(selection[0])
                voice_path = str(Path(folder) / voice_name)
                
                output = Path("_test_audio.wav")
                
                self.status_label.config(text="Generating test audio...")
                self.model.tts_to_file(
                    text=text,
                    file_path=str(output),
                    speaker_wav=voice_path,
                    language=self.language_var.get(),
                    speed=0.9
                )
                
                self.status_label.config(text="Playing test audio...")
                pygame.mixer.music.load(str(output))
                pygame.mixer.music.play()
                
                # Clean up after playing
                while pygame.mixer.music.get_busy():
                    time.sleep(0.1)
                output.unlink(missing_ok=True)
                
                self.status_label.config(text="Test completed successfully")
                
            except Exception as e:
                self.status_label.config(text=f"Error: {str(e)}")
                messagebox.showerror("Error", f"Failed to generate test audio:\n{e}")
                
        threading.Thread(target=generate, daemon=True).start()
        
    def generate_audio(self):
        """Generate audio files from document."""
        if self.is_generating:
            messagebox.showinfo("Busy", "Already generating audio files.")
            return
            
        selection = self.voice_listbox.curselection()
        if not selection:
            messagebox.showwarning("No Voice", "Please select a voice first.")
            return
            
        if not self.selected_docx:
            messagebox.showwarning("No Document", "Please select a document first.")
            return
            
        def process():
            self.is_generating = True
            self.generate_btn.config(state=DISABLED, text="⏳ Generating...")
            
            try:
                self.load_tts_model()
                
                folder = self.voice_folder_entry.get()
                voice_name = self.voice_listbox.get(selection[0])
                voice_path = str(Path(folder) / voice_name)
                
                # Read document
                doc = Document(self.selected_docx)
                paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                
                # Create output folder
                doc_name = sanitize_name(Path(self.selected_docx).stem)
                output_folder = Path(self.output_dir) / f"{self.language_var.get()}_{doc_name}"
                output_folder.mkdir(parents=True, exist_ok=True)
                
                # Generate audio files
                total = len(paragraphs)
                errors = []
                
                for i, text in enumerate(paragraphs):
                    self.status_label.config(text=f"Processing {i+1}/{total}...")
                    self.progress['value'] = (i + 1) / total * 100
                    self.root.update()
                    
                    output_file = output_folder / f"audio_{i+1}.wav"
                    
                    try:
                        self.model.tts_to_file(
                            text=text,
                            file_path=str(output_file),
                            speaker_wav=voice_path,
                            language=self.language_var.get(),
                            speed=0.85
                        )
                    except Exception as e:
                        errors.append((i+1, text, str(e)))
                        continue
                
                # Save errors if any
                if errors:
                    error_doc = Document()
                    error_doc.add_heading("Paragraphs with Errors", 0)
                    for idx, text, error in errors:
                        error_doc.add_heading(f"Paragraph {idx}", 1)
                        error_doc.add_paragraph(text)
                        error_doc.add_paragraph(f"Error: {error}")
                    error_doc.save(output_folder / "errors.docx")
                
                self.progress['value'] = 100
                self.status_label.config(text=f"Completed! {total - len(errors)}/{total} files generated")
                
                messagebox.showinfo(
                    "Success",
                    f"Generated {total - len(errors)} audio files!\n\nOutput folder:\n{output_folder}"
                )
                
                # Open output folder
                os.startfile(output_folder) if os.name == 'nt' else os.system(f'open "{output_folder}"')
                
            except Exception as e:
                self.status_label.config(text=f"Error: {str(e)}")
                messagebox.showerror("Error", f"Failed to generate audio:\n{e}")
                
            finally:
                self.is_generating = False
                self.generate_btn.config(state=NORMAL, text="⚡ Generate All Audio Files")
                self.progress['value'] = 0
                
        threading.Thread(target=process, daemon=True).start()

# =============================
# MAIN ENTRY POINT
# =============================
def main():
    root = Tk()
    app = ModernTTSApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()